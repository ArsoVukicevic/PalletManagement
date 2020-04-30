# Namespace - Biblioteka pojedinacnih funkcionalnosti aplikacije
#Lista funkcionalnosti 
    #1 - Generisanje QR koda(images) 
    #2 - Stampanje images na default stampacu na Windows OS-u
    #3 - Citanje QR koda sa images + odredjivanje pozicije (bounding box)
    #4 - Akvizicija images sa USB/IP Cameras
    #5 - Generisanje panorame na osnovu vise slika (image stitching)
import matplotlib.pyplot as plt
import numpy as np
import cv2 as cv 
import imutils
from PyQt5 import Qt

#4 - Acquisition of images from USB/IP Cameras
#ref:
#INPUTS
    # Cameras - list of cameras (USB id, or IP cameras settings)
#OUTPUTS
    # rezImg - list of images
    # status - strin with info is camera available
def takeImageFromCamera(Cameras, debugMode=False):
    import cv2 as cv 
    import imutils
    import numpy as np
    rezImg = []
    status = []
    for iCamera in Cameras:
        cap = cv.VideoCapture(iCamera)
        return_value, image = cap.read()
        if return_value:
            rezImg.append(image)
            status.append(str(iCamera))
            if debugMode:
                while True:
                    return_value, image = cap.read()
                    cv.imshow('img1', image)
                    k = cv.waitKey(100) 
                    if k == 27: #press Esc to exit
                        break
        else:
            rezImg.append(np.zeros((512,512, 3), np.uint8))  
            status.append(str(iCamera) + ' nedostupna')
        cap.release() # 
        cv.destroyAllWindows()  
    return rezImg, status
# END take kameru sa usb/ip Cameras -----------------------------------------------------


def createPanorama(imgs,debugMode=False):
    import cv2 as cv
    stitcher     = cv.Stitcher.create(cv.Stitcher_PANORAMA)#Stitcher_SCANS
    status, pano = stitcher.stitch(imgs)
    if type(pano) == type(None):
        pano = []
    if debugMode and pano != []:
        cv.imshow('img',pano)  
        cv.waitKey(0)
        cv.destroyAllWindows()
    return pano    
    #cv.imwrite("test imgs panorama/rez.jpg", pano)
# # ---- END Panorama sitching ----------------------------------------------------------------------


#2 - Stampanje images na default stampacu na Windows OS-u
#ref: https://gist.github.com/buptxge/2fc61a3f914645cf8ae2c9a258ca06c9
#INPUT
    #inputImg - slika QR koda koja treba da se stampa u NumPy array formatu
def printQqCodeOnWinPrinter(inputImg, debugMode):
    import win32print
    import win32ui
    from PIL import Image, ImageWin
    # Constants for GetDeviceCaps
    # HORZRES / VERTRES = printable area
    HORZRES         = 8
    VERTRES         = 10
    # LOGPIXELS = dots per inch
    LOGPIXELSX      = 300
    LOGPIXELSY      = 600
    # PHYSICALWIDTH/HEIGHT = total area
    PHYSICALWIDTH   = 110
    PHYSICALHEIGHT  = 111
    # PHYSICALOFFSETX/Y = left / top margin
    PHYSICALOFFSETX = 112
    PHYSICALOFFSETY = 113

    printer_name    = win32print.GetDefaultPrinter ()
    file_name       = "arsQR.png"
    #bmp = Image.open(file_name)
    hDC = win32ui.CreateDC ()
    hDC.CreatePrinterDC(printer_name)
    printable_area  = hDC.GetDeviceCaps (HORZRES), hDC.GetDeviceCaps (VERTRES)
    printer_size    = hDC.GetDeviceCaps (PHYSICALWIDTH), hDC.GetDeviceCaps (PHYSICALHEIGHT)
    printer_margins = hDC.GetDeviceCaps (PHYSICALOFFSETX), hDC.GetDeviceCaps (PHYSICALOFFSETY)
    
    bmp = Image.fromarray(inputImg)
    if bmp.size[0] > bmp.size[1]:
        bmp = bmp.rotate(90)

    ratios = [1.0 * printable_area[0] / bmp.size[0], 1.0 * printable_area[1] / bmp.size[1]]
    scale = min (ratios)

    hDC.StartDoc(file_name)
    hDC.StartPage()
    printable_area = hDC.GetDeviceCaps (HORZRES), hDC.GetDeviceCaps (VERTRES)

    dib = ImageWin.Dib(bmp)
    scaled_width, scaled_height = [int (scale * i) for i in bmp.size]
    x1  = int ((printer_size[0] - scaled_width) / 2)
    y1  = int ((printer_size[1] - scaled_height) / 2)
    x2  = x1 + scaled_width
    y2  = y1 + scaled_height

    dib.draw(hDC.GetHandleOutput (), (x1, y1, x2, y2))
    hDC.EndPage ()
    hDC.EndDoc ()
    hDC.DeleteDC ()
#end printerv ----------------------------------------------------------------------


#1 - Generate QR code(image)  ---------------------------------------------------------------------
#ref: https://github.com/lincolnloop/python-qrcode
#INPUTS
    #stringData        - string that will be written on QR code
    #debugMode         - true=draw, false=write
#OUTPTUS
    #rezQRimg          - generisana slika
def generateQR(stringData, debugMode=False):
    import qrcode # modul koji treba instalirati https://github.com/lincolnloop/python-qrcode
    qr = qrcode.QRCode(
        version          = 2, 
        error_correction = qrcode.constants.ERROR_CORRECT_L,
        box_size         = 11, #10 default
        border           = 4, #4 default
    )
    qr.add_data(stringData)
    qr.make(fit=True)
    # Pravljanje images od strukture qrcode biblioteke
    img       = qr.make_image(fill_color="black", back_color="white")
    rezQRimg  = img._img.convert('RGB')
    rezQRimg  = rezQRimg.crop((36,36, 330,330))
    # Dodaj tekst na dnu images
    sirinaQR, visinaQR = rezQRimg.size
    from PIL import Image, ImageDraw, ImageFont
    img  = Image.new('RGB', (sirinaQR+17, visinaQR+55), color = 'yellow')
    img.paste(rezQRimg, (5, 5, 5+sirinaQR, 5+visinaQR))
    draw = ImageDraw.Draw(img)
    draw.text((10,visinaQR+10), 'QR kod: ' + stringData, 'rgb(0, 0, 0)', ImageFont.truetype('arial.ttf', size=20))
    import datetime
    currentTime = str(datetime.datetime.now())
    draw.text((10,visinaQR+30), 'currentTime : ' + currentTime[:19], 'rgb(0, 0, 0)', ImageFont.truetype('arial.ttf', size=20))
    rezQRimg = img.convert('RGB')
    #Ukoliko treba stampati medjurezultate
    if debugMode:
        rezQRimg.save('_debug_rezQRimg.png')
        import  matplotlib.pyplot as plt
        plt.figure(facecolor='white')
        plt.text(0.5, 0.5, stringData, horizontalalignment='left', fontsize=16)
        plt.imshow(rezQRimg)
        plt.axis('off')
        plt.savefig('_debug_rezQRimg_PyPlot.png')
        plt.show()    
    import numpy    
    return numpy.array(rezQRimg.convert('RGB')) #konvertuje se u format citljiv OpenCV i dr bibliotekama   
# end pisanje QRkoda ------------------------------------------------------------------

# 3 - Read QR code from image and return bounding boxes--------------------------------
#ref: https://pypi.org/project/pyzbar/
#INPUTS
    # inputImg - input image with mutiple QR codes
#OUTUTS
    # data         - list of strings
    # dataLocaions - list ob bboxes
def readQRcode(inputImg, debugMode=False):
    #Ukljuci potrebne module https://pypi.org/project/pyzbar/
    from pyzbar.pyzbar import decode
    from PIL import Image
    rez                    = decode(inputImg)
    numOfDetectedQRs        = len(rez)
    data                   = []
    dataLocaions           = []
    for i in range(numOfDetectedQRs):
            data.append(rez[i].data.decode('UTF-8'))
            dataLocaions.append(rez[i].polygon)
    if debugMode:
        print("Number of qr codes is: ", str(numOfDetectedQRs))
        print("Read QR codes are:")
        import  matplotlib.pyplot as plt
        plt.figure(facecolor='white')
        plt.imshow(inputImg)        
        for i in range(numOfDetectedQRs):
            print("Code#", str(i+1),":", str(rez[i].data))
            plt.text(rez[i].polygon[0].x, rez[i].polygon[0].y, rez[i].data, horizontalalignment='left', fontsize=16)
            rectangle = plt.Polygon(rez[i].polygon,closed=True, fill=None, edgecolor='r')
            plt.gca().add_patch(rectangle)
        print(type(rez))
        plt.axis('off')
        plt.show()  
    return data, dataLocaions
#end citanje QRkoda ---------------------------------------------------------------------

# Fja za nalazenje objekta na slici. 
#ref: https://opencv-python-tutroals.readthedocs.io/en/latest/py_tutorials/py_feature2d/py_feature_homography/py_feature_homography.html
#ref: http://amroamroamro.github.io/mexopencv/opencv/feature_homography_demo.html 
#INPUTS
    #imgScene - global image being searched
    #imgObj   - objekat koji se trazi na slici
def nadjiObjekatNaSlici(imgScene, imgObj, debugMode=False):
    import cv2 as cv
    import numpy as np
    detector             = cv.ORB_create() 
    [keyObj, featObj]    = detector.detectAndCompute(imgObj, None) # compute the descriptors with ORB
    [keyScene,featScene] = detector.detectAndCompute(imgScene, None)
    FLANN_INDEX_LSH=6
    index_params= dict(algorithm     = FLANN_INDEX_LSH,
                   table_number      = 6, # 12
                   key_size          = 12,     # 20
                   multi_probe_level = 1) #2
    search_params = dict(checks=50)   # or pass empty dictionary
    flann         = cv.FlannBasedMatcher(index_params,search_params)
    matches       = flann.knnMatch(featObj,featScene,k=2)
    # Apply ratio test
    good = []
    for i in range(len(matches)):
        try:        
            if matches[i][0].distance < 0.75*matches[i][1].distance:
                    good.append(matches[i][0])
        except:
            a=1
    img_match = np.empty((max(imgObj.shape[0], imgScene.shape[0]), imgObj.shape[1] + imgScene.shape[1], 3), dtype=np.uint8)
    img3 = cv.drawMatches(imgObj,keyObj,imgScene,keyScene,good,None,flags=2)
    cv.imshow("flann matching", img3)
    cv.waitKey(0) 
    # Homografija
    ptsObj   = []
    ptsScene = []
    ptsObj   = np.float32([keyObj[m.queryIdx].pt for m in good]).reshape(-1,1,2)
    ptsScene = np.float32([keyScene[m.trainIdx].pt for m in good]).reshape(-1,1,2)

    [M,mask] = cv.findHomography(ptsObj, ptsScene, cv.RANSAC,5.0)
    matchesMask = mask.ravel().tolist()
    h,w, bzv = imgObj.shape
    pts = np.float32([ [0,0],[0,h-1],[w-1,h-1],[w-1,0] ]).reshape(-1,1,2)
    dst = cv.perspectiveTransform(pts,M)

    img2 = cv.polylines(imgScene,[np.int32(dst)],True,255,3, cv.LINE_AA)
    cv.imshow("flann matching", img2)
    cv.waitKey(0) 
    return M, dst # vraca se matrica Homografije i bounidng box images koja se trazi

# ref: https://python-docx.readthedocs.io/en/latest/#contributor-guide
def PrintMsWordReport(ListaTaba, Panorama=[]):
    ListaTaba = odradiAnalizuSkladista(ListaTaba)
    from docx import Document
    from docx.shared import Inches
    import datetime
    document = Document()
    time     = datetime.datetime.now()
    document.add_heading('Warehouse inventory' , 0)
    p1       = document.add_paragraph('Time of the inventory check: ' + str(time))
    document.add_heading('List of detected tabs' + '(' + str(len(ListaTaba)) + ')' , 0)
    WO       = -1
    for i in range(len(ListaTaba)):
        if WO != int(ListaTaba[i][:4]):
            WO = int(ListaTaba[i][:4])
            p2 = document.add_paragraph('Work orders' + str(WO), style='List Bullet')
        p2 = document.add_paragraph('\t Tab ' + str(i+1) + ': ' + ListaTaba[i])        
    document.add_page_break()
    document.save('GrafoStill_inventory.docx')
    #subprocess.Popen(["C:\\Program Files\Microsoft Office\Office12\winword.exe", "demo.docx", "/mFilePrintDefault", "/mFileExit"]).communicate()

# Fja pravi listu kompletnih i nekompletnih Taba
#INPUTS
    #ListaTaba - lista Taba koja je pokupljena sa IP kamera
#OUTPUTS
    #ListaTabaPopis
def odradiAnalizuSkladista(ListaTaba):
    #napravi listu WO (radnih naloga) u skladistu
    rez = []
    for Tab in ListaTaba:
        rez.append(Tab.qrData)     
    rez.sort()
    return rez
#-------------------------------------------------------------

# Klasa IpCamera je definisana svojom IP/USB adresom, na osnvou koje konstruktor cita ostale podatke
class IpCamera ():
    adresa      = [] # adresa IP Cameras ili id USB Cameras
    status      = [] # string, camera is available or not
    img         = [] # images from Cameras
    qrData      = [] # QR data
    qrLocations = [] # location of detected QR codes
    def __init__(self, ip=[]):
        self.setAdress(ip)
        self.refreshData()
    # IP aresa je vec postavljena - ocitaj Image sa Cameras; ocitaj QR kodove sa images
    def refreshData(self):        
        self.getImageFromCamera()
        self.readQRkod()
    def setAdress(self, Adresa=[]):
        self.adresa = Adresa
    def getImageFromCamera(self):
        img, status = takeImageFromCamera([self.adresa])
        self.img    = img[0]
        self.status = status[0]
    def readQRkod(self):
        self.qrData, self.qrLocations = readQRcode(self.img)
    def getDetectedQrTabs(self):
        qrTabs = []
        for i in range(len(self.qrData)):
            qrTabs.append(QrTab(self.qrData[i], self.qrLocations[i], self.img, self.adresa, i))
        return qrTabs
        
# Klasa QrTab u sebi cuva podatke o Tabu koji je opisan QR kodom koji je detektovan na nekoj slici
class QrTab():
    qrData     = [] # 0001,01/01'
    qrLocation = [] # bbox
    qrImage    = [] # 
    qrCamera   = [] # id/ip adresss
    idQrCode   = [] # id of the code on qrCamera

    def __init__(self, data=[], Location=[], Image=[], Camera=[], QrCode=[]):
        self.qrData     = data
        self.qrLocation = Location
        self.qrImage    = Image
        self.qrCamera   = Camera
        self.idQrCode   = QrCode

    def checkQrDataFormat(self, qrData=[]):
        try: 
            IdRadnogNaloga, IdTaba, UkupnoTaba = self.parsujqrData(qrData)
            procitaniString = '{:04d},{:02d}/{:02d}'.format(IdRadnogNaloga, IdTaba, UkupnoTaba)
            if procitaniString == self.qrData: #dodatna provera da li je sam format zapisa validan
                if IdTaba <= UkupnoTaba:
                    return True  # Sve ok
                else:
                    return False # Id Taba mora biti manji od ukupno Taba
            else:
                return False     # Mozda je ubacio neki string 0001,01/0A npr
        except:
            return False         # QR kod ne moze da se procita

    def parsujqrData(self, qrKod = []):
        if qrKod == []:
            qrKod = self.qrData
        IdRadnogNaloga = int(qrKod[:4])
        IdTaba = int(qrKod[5:7])
        UkupnoTaba = int(qrKod[8:10])
        return IdRadnogNaloga, IdTaba, UkupnoTaba

# Work order - cuva listu Taba. Moze da proveri da li su svi an broju, regulaWOo rasporedjeni itd.
class WorkOrder():
    listaTaba            = [] # lista Taba koji pripadaju WO
    listaTabaSaQrGreskom = [] # posebna lista za Tabe sa greskom
    ukupanBrTaba         = 0  # ukupan br Taba u WO (odredjen je QR kodom prvog Taba)
    idRadnogNaloga         = 0  # id radnog naloga (odredjen je QR kodom prvog Taba)
    # konstruktor
    def __init__(self, tabaci):
        self.listaTaba            = []
        self.listaTabaSaQrGreskom = []
        if type(tabaci) == QrTab:
            tabaci = [tabaci]
        for Tab in tabaci:
            self.dodajTab(Tab)
        self.sortirajTabe()
        
    def dodajTab(self, Tab):
        #ID radnog naloga i ukupan br Taba odredjeni su prvim Tabom koji se dodaje
        idRadnogNaloga, idTaba, ukupnoTaba = Tab.parsujqrData()
        if self.listaTaba == []:            
            self.ukupanBrTaba = ukupnoTaba
            self.idRadnogNaloga = idRadnogNaloga
        # da li Tab pripada ovom WO           
        if  self.idRadnogNaloga == idRadnogNaloga:                                 # ako pripada ovom WO 
            if (ukupnoTaba <= self.ukupanBrTaba) and (idTaba <= ukupnoTaba):# ako je QR kod validan
                if any(x.qrData == Tab.qrData for x in self.listaTaba):  # ako vec postoji u listiTaba 
                    pass                                                           # ne radi nista
                else:                                                              # ako je razlicit od postojecih QR kodova
                    self.listaTaba.append(Tab)                                   # dodaj u listuTaba
            else:                                                                  # ako QR kod nije validan
                self.listaTabaSaQrGreskom.append(Tab)                            # dodaj u listaTabaSaQrGreskom

    # Fja proverava da li je QR/Tab detektovan i da li je validan
    # INPUTS
        # qrKod - string
    # OUTPUTS
        # Tab - objekat QrTab
        # da li je validan - bool    
    def getTab(self, qrKod):
        rez = []
        idRadnogNaloga, idTaba, ukupnoTaba = QrTab.parsujqrData(qrKod)
        for Tab in self.listaTaba:
            if Tab.qrData == qrKod:
                return Tab, True
        for Tab in self.listaTabaSaQrGreskom:
            if Tab.qrData == qrkod:
                return Tab, False

   # Fja vraca string-listu sa tabacima koji imaju gresku u QR kodu
    #INPUTS
        # self.tabaciSaQrGreskom
    #OUTPUTS
        # stringListaTabaSaQrGreskom - string lista QR kodova
        # self.tabaciSaQrGreskom       - lista objekata (QrTab)
    def getListuTabaSaQrGreskomKaoString(self):
        stringListaTabaSaQrGreskom = []
        for Tab in self.tabaciSaQrGreskom:
            stringListaTabaSaQrGreskom.append(Tab.qrData)
        return stringListaTabaSaQrGreskom, self.tabaciSaQrGreskom

    # Fja generise string-litu detektovanih Taba 
    #INPUTS
        # self.tabaci - lista detektovanih Taba
    #OUTPUTS
        # stringListTabsWO        - string lista QR kodova
        # stringListTasValidity - Validity Taba
        # self.tabaci                - lista objekata (QrTab)    
    def getListTabsAsString(self):
        stringListTabsWO        = []
        stringListTasValidity = []
        rezListTabs             = []
        for Tab in self.listaTaba:
            stringListTabsWO.append(Tab.qrData)
            rezListTabs.append(Tab)
        for Tab in self.listaTabaSaQrGreskom:
            stringListTabsWO.append(Tab.qrData)
            rezListTabs.append(Tab)
        if stringListTabsWO == []:
            return [], [], []
        stringListTabsWO.sort(key=lambda x: x, reverse=False)
        pomTab = QrTab(rezListTabs[0].qrData, rezListTabs[0].qrLocation, rezListTabs[0].qrImage, rezListTabs[0].qrCamera,  rezListTabs[0].idQrCode)
        for qrData in stringListTabsWO:            
            pomTab.qrData =  qrData
            stringListTasValidity.append(pomTab.checkQrDataFormat(qrData))
        return stringListTabsWO, stringListTasValidity, rezListTabs
    
    # Fja generise string-litu Taba koji nedostaju
    #INPUTS
        # None
    #OUTPUTS
        # listaNedostajucihTabaWO - string lista QR kodova koji nedostaju
    def getListuNedostajucihTabaKaoString(self):
        listaNedostajucihTabaWO = []  
        for iTaba in range(self.ukupanBrTaba):
            qrKod = '{:04d},{:02d}/{:02d}'.format(self.idRadnogNaloga, iTaba, self.ukupanBrTaba)
            Tab = self.getTab(qrKod)
            if Tab == []:
                listaNedostajucihTabaWO.append(qrKod)
        return listaNedostajucihTabaWO  

    def sortirajTabe(self):
        self.listaTaba.sort(key=lambda x: x.qrData, reverse=False)

    def CheckIfAllTabsArePresent(self):
        return  len(self.listaTaba) == self.ukupanBrTaba

    #OUTPUTS
        # Validity: TRUE - sve ok    FALSE - ima nepravilno oznacenih QR kodov
    def checkIfAllTabsAreValid(self):
        return len(self.listaTabaSaQrGreskom) ==  0

    # ref: https://python-docx.readthedocs.io/en/latest/#contributor-guide
    def PrintMsWordReport(self):
        statusWordOrder_text          = dict([ ('TrueTrue', 'valide and complete.'), ('TrueFalse', 'valide and incomplete.'), ('FalseTrue', 'invalid and complete.'), ('FalseFalse', 'invalid i incomplete.') ])
        statusWordOrder_text_Validity = dict([(True, '.'), (False, ' is invalid.')])
        from docx import Document
        from docx.shared import Inches
        import datetime
        document       = Document()
        currentTime    = str(datetime.datetime.now())
        currentTime    =  currentTime[:13] + 'h ' + currentTime[14:16] + 'm ' + currentTime[17:19] + 's'
        document.add_heading('GRAFOSTIL Kragujevac' , 0)
        p1 = document.add_paragraph('Inventory time: ' + currentTime)
        document.add_heading('Work order ' + '{:04d}'.format(self.idRadnogNaloga) , 0)
        WO_validity     = self.checkIfAllTabsAreValid()
        WO_completness  = self.CheckIfAllTabsArePresent() 
        p1 = document.add_paragraph('Work order is ' + statusWordOrder_text[str(WO_validity) + str(WO_completness)])
        # Podaci po Radnim nalozima   
        stringListTabsWO, stringListTasValidity, rezListTabs = self.getListTabsAsString()
        document.add_heading('List of detected Tabs' + '(' + str(len(stringListTabsWO)) + '/' + str(self.ukupanBrTaba)   + ')' , 0)
        for (stringTab, stringValidity) in zip(stringListTabsWO, stringListTasValidity):
            if stringValidity:
                p2 = document.add_paragraph('\t Tab ' + stringTab)
            else:
                p2 = document.add_paragraph('\t Tab ' + stringTab + ' is invalid !!!' )
        document.add_page_break()
        document.save('GrafoStill - Report for  WO ' + '{:04d}'.format(self.idRadnogNaloga) + ' ' + currentTime + '.docx')
    def proveriDaLiSuTabaciPravilnoRasporedjeni():
        a=1 #potrebna panorama

class Warehouse():
    ipCameras        = []
    detectedTabs     = []
    workOrders       = []
    panorama         = []

    def __init__(self, ipAdrese = [0,1,2,3]):
        for ip in ipAdrese:
            self.ipCameras.append(IpCamera(ip))
        self.refreshDataFromIpCameras()
        
    def refreshDataFromIpCameras(self):
        self.detectedTabs = []
        self.workOrders       = []
        for iCameras in range(len(self.ipCameras)):
            self.ipCameras[iCameras].refreshData()
            for Tab in self.ipCameras[iCameras].getDetectedQrTabs():
                self.detectedTabs.append(Tab)
        # ako nije nadjen nijedan Tab prekini dalje izvrsavanje fje
        workOrders        = []
        imgsZaPanoramu    = []
        if len(self.detectedTabs) == 0:
            return
        # sortiraj detektovane Tabe      
        self.detectedTabs.sort(key=lambda x: x.qrData, reverse=False)    
        # izbaci duplikate - tabaci su duplikati ako imaju isti QR kod
        jedinstveniTabaci = [self.detectedTabs[0]]
        for Tab in self.detectedTabs:
            if any(x.qrData == Tab.qrData for x in jedinstveniTabaci): # vec postoji
                continue 
            else:        # nov
                jedinstveniTabaci.append(Tab)
        self.detectedTabs = jedinstveniTabaci
        # #razvrstaj Tabe na WO (ako  neki Tab ima nevalidan QR kod-to se resava na nivou WO)
        for Tab in self.detectedTabs:
            idRadnogNaloga, idTaba, ukupnoTaba = Tab.parsujqrData()
            if any(x.idRadnogNaloga == idRadnogNaloga for x in self.workOrders): # ako vec postoji Work order kome Tab pripada
                self.workOrders[-1].dodajTab(Tab) 
            else:
                self.workOrders.append(WorkOrder(Tab))
        for iCamera in self.ipCameras:
            if iCamera.qrData != []:
                imgsZaPanoramu.append(iCamera.img)
        #panorama = createPanorama(imgsZaPanoramu, True) # - trebapodesiti kamere da se preklapaju !!!
    

    def getImagesFromIpCameras(self):
        images  = []
        status = []
        for kamera in self.ipCameras:
            images.append(kamera.img)
            status.append(kamera.status)
        return images, status

    def getStringListuQrData(self):
        qrData = []
        qrDataIsValid = []
        for Tab in self.detectedTabs:
            qrData.append(Tab.qrData)
            qrDataIsValid.append(Tab.checkQrDataFormat())
        return qrData, qrDataIsValid
    
    def checkIfThereIsAlreadySuchQrCodeOnInventory(self,QrKod):
        for Tab in self.detectedTabs:
            if Tab.qrData == QrKod:
                return True, Tab
        return False, []
    
    def getStringListOfWorkOrders(self):
        stringListaWO             = []
        boolListaWO_validityWO   = []
        boolListaWO_completnessWO = []
        for WO in self.workOrders:
            stringListaWO.append('{:04d}'.format(WO.idRadnogNaloga))        
            boolListaWO_validityWO.append(WO.checkIfAllTabsAreValid())
            boolListaWO_completnessWO.append(WO.CheckIfAllTabsArePresent())
        return stringListaWO, boolListaWO_validityWO, boolListaWO_completnessWO

    # ref: https://python-docx.readthedocs.io/en/latest/#contributor-guide
    def PrintMsWordReport(self):
        statusWordOrder_text = dict([ ('TrueTrue', 'valid and complete.'), ('TrueFalse', 'valid and incomplete.'), ('FalseTrue', 'invalid and complete.'), ('FalseFalse', 'invalid and incomplete.') ])
        statusWordOrder_text_Validity = dict([(True, ''), (False, ' is invalid!!!')])
        #
        from docx import Document
        from docx.shared import Inches
        import datetime
        document = Document()
        currentTime = str(datetime.datetime.now())
        currentTime = currentTime[:13] + 'h ' + currentTime[14:16] + 'm ' + currentTime[17:19] + 's'
        document.add_heading('GRAFOSTIL Kragujevac' , 0)
        p1 = document.add_paragraph('Time of inventory: ' + currentTime)
        document.add_heading('List of detected work orders' + '(' + str(len(self.workOrders)) + ')' , 0)
        #
        stringListaWO, boolListaWO_validityWO, boolListaWO_completnessWO = self.getStringListOfWorkOrders()
        p1 = document.add_paragraph('Complete work orders (' + str(sum(boolListaWO_completnessWO))  + ')', style='List Bullet')
        for (WO, WO_validity, WO_completness) in zip(stringListaWO, boolListaWO_validityWO, boolListaWO_completnessWO):
            if WO_completness == True:            
                p2 = document.add_paragraph('\t WO ' + WO + statusWordOrder_text_Validity[WO_validity])
        #
        p1 = document.add_paragraph('Inomplete work orders(' + str(len(boolListaWO_completnessWO)-sum(boolListaWO_completnessWO)) + ')', style='List Bullet')
        for (WO, WO_validity, WO_completness) in zip(stringListaWO, boolListaWO_validityWO, boolListaWO_completnessWO):
            if WO_completness == False:            
                p2 = document.add_paragraph('\t WO ' + WO + statusWordOrder_text_Validity[WO_validity])
        # Podaci po Radnim nalozima   
        document.add_heading('Detected Tabs' + '(' + str(len(self.detectedTabs)) + ')', 0)
        for WO in self.workOrders:
            statusWO = str(WO.checkIfAllTabsAreValid()) + str(WO.CheckIfAllTabsArePresent())    
            stringListTabsWO, stringListTasValidity, rezListTabs = WO.getListTabsAsString()        
            p2 = document.add_paragraph('WO ' + '{:04d}'.format(WO.idRadnogNaloga) + ' je ' + statusWordOrder_text[statusWO] + ' Detected Tabs (' + str(len(stringListTabsWO)) + '/' + str(WO.ukupanBrTaba) + ').', style='List Bullet')
            for (stringTab, stringValidity) in zip(stringListTabsWO, stringListTasValidity):
                if stringValidity:
                    p2 = document.add_paragraph('\t Tab ' + stringTab )
                else:
                    p2 = document.add_paragraph('\t Tab ' + stringTab + ' is invalid !!!' )
        document.add_page_break()
        document.save('GrafoStill - Warehouse inventory' + currentTime + '.docx')
  

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas 

class PlotCanvas(FigureCanvas):
    def __init__(self, parent=None, width=5, height=4, dpi=100):
        from matplotlib.figure import Figure
        fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = fig.add_subplot(111)
        FigureCanvas.__init__(self, fig)
        self.setParent(parent)
        from PyQt5.QtWidgets import QSizePolicy
        FigureCanvas.setSizePolicy(self,
                QSizePolicy.Expanding,
                QSizePolicy.Expanding)
        FigureCanvas.updateGeometry(self)
        self.drawImage()
    
    def drawImage(self, img=[], qrLocations=[], qrData=[]):
        figSize = [self.size().height(), self.size().width()]
        self.figure.clear()
        if img == []:
            import numpy as np
            img.append(np.zeros((figSize[0], figSize[1], 3), np.uint8))
        import matplotlib.pyplot as plt
        self.figure, ax = plt.subplots() 
        self.figure.set_figwidth(figSize[1]/100)
        self.figure.set_figheight(figSize[0]/100)
        plt.imshow(img[0], interpolation='none')
        plt.axis('off') 
        plt.subplots_adjust(left=0, right=1, top=1, bottom=0)  
        ax.set_aspect('auto')
        if qrLocations != []:
            for i in range(len(qrLocations)):
                rectangle = plt.Polygon(qrLocations[i],closed=True, fill=None, edgecolor='r')
                plt.gca().add_patch(rectangle)
                plt.text(qrLocations[i][0].x, qrLocations[i][0].y, qrData[i], horizontalalignment='left', fontsize=16, backgroundcolor='white')
        plt.close()
        self.draw()